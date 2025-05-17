# PythonEngine/output_writer.py

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

def save_output(text_result, output_txt_path):
    # Simpan hasil teks ke file txt
    with open(output_txt_path, 'w', encoding='utf-8') as f:
        f.write(text_result)

    # Parsing teks spesifik untuk Excel (opsional)
    parsed_rows = []
    current = {}

    for line in text_result.splitlines():
        line = line.strip()
        if line.startswith("Masalah:"):
            current['Masalah'] = line.replace("Masalah:", "").strip()
        elif line.startswith("Rekomendasi:"):
            current['Rekomendasi'] = line.replace("Rekomendasi:", "").strip()
        elif line.startswith("Catatan:"):
            current['Catatan'] = line.replace("Catatan:", "").strip()
            parsed_rows.append(current)
            current = {}

    if parsed_rows:
        df = pd.DataFrame(parsed_rows)
        output_excel_path = output_txt_path.replace("hasil_output.txt", "hasil_parsed.xlsx")
        df.to_excel(output_excel_path, index=False)

def save_as_txt(text: str, path: str):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"[SUKSES] Teks disimpan ke {path}")

def save_as_word(text: str, path: str):
    doc = Document()
    
    # Setup dokumen dengan style yang lebih baik
    # 1. Set margin dokumen
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 2. Style untuk heading
    heading_style = doc.styles['Heading 1']
    heading_font = heading_style.font
    heading_font.name = 'Calibri'
    heading_font.size = Pt(16)
    heading_font.bold = True
    heading_font.color.rgb = RGBColor(0, 0, 0)
    
    # 3. Style untuk normal paragraf
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 4. Buat style untuk subheading
    if 'Subheading' not in doc.styles:
        subheading_style = doc.styles.add_style('Subheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_font = subheading_style.font
        subheading_font.name = 'Calibri'
        subheading_font.size = Pt(14)
        subheading_font.bold = True
        subheading_font.italic = True
    
    # Parse markdown header indicators
    lines = text.strip().splitlines()
    
    # Deteksi jenis konten
    is_markdown_table = any('|' in line and '---' in line for line in lines[:10])
    is_biography = 'biografi' in text.lower() or 'profil' in text.lower()
    has_section_headers = any(line.startswith('#') for line in lines)
    
    if is_markdown_table and not is_biography:
        # Tambahkan header jika ada tabel tapi tidak eksplisit ada header
        if not has_section_headers:
            header = doc.add_heading('Hasil Data', level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Konversi tabel markdown ke tabel Word
        table_started = False
        table_lines = []
        
        for line in lines:
            if '|' in line and not table_started:
                table_started = True
            
            if table_started and '|' in line:
                table_lines.append(line)
                
        if table_lines:
            headers = [cell.strip() for cell in table_lines[0].strip('|').split('|')]
            data_rows = []
            for line in table_lines[2:]:  # Skip separator line
                if '|' in line:
                    data_rows.append([cell.strip() for cell in line.strip('|').split('|')])
            
            # Buat tabel Word
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Tambahkan header
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.bold = True
            
            # Tambahkan data rows
            for row_data in data_rows:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    if i < len(row_cells):  # Pastikan tidak melebihi jumlah kolom
                        if cell_text.startswith('='):
                            row_cells[i].text = f"Formula: {cell_text}"
                        else:
                            row_cells[i].text = cell_text
    
    elif is_biography:
        # Format khusus untuk biografi
        # Tambahkan judul
        title = "Biografi"
        for line in lines[:5]:
            if "biografi" in line.lower() or "profil" in line.lower():
                title = line.strip()
                break
        
        header = doc.add_heading(title, level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse konten biografi
        if is_markdown_table:
            # Jika biografi dalam format tabel, konversi ke paragraf terstruktur
            for line in lines[2:]:  # Skip header + separator
                parts = [cell.strip() for cell in line.strip('|').split('|')]
                if len(parts) >= 2:
                    p = doc.add_paragraph()
                    p.add_run(f"{parts[0]}: ").bold = True
                    p.add_run(parts[1])
        else:
            # Format biografi dalam bentuk paragraf
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Deteksi header
                if line.startswith('#'):
                    level = line.count('#')
                    title = line.strip('#').strip()
                    if level == 1:
                        doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        doc.add_heading(title, level=min(level, 3))
                    current_section = title.lower()
                
                # Deteksi point seperti "Nama: John Doe"
                elif ': ' in line and len(line.split(': ')[0]) < 30:
                    key, value = line.split(': ', 1)
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(value)
                
                # List items
                elif line.startswith('* ') or line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                
                # Paragraf biasa
                else:
                    # Cek apakah ini bisa menjadi subheading
                    if len(line) < 50 and line.isupper() or line.endswith(':'):
                        p = doc.add_heading(line, level=2)
                        current_section = line.lower().replace(':', '')
                    else:
                        p = doc.add_paragraph(line)
                        # Jika dalam seksi biografi/latar belakang, justify teksnya
                        if current_section and ('latar' in current_section or 'biografi' in current_section):
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    else:
        # Format untuk dokumen umum
        current_level = 0
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Deteksi markdown headers
            if line.startswith('#'):
                level = min(line.count('#'), 3)  # Max level 3
                title = line.strip('#').strip()
                h = doc.add_heading(title, level=level)
                if level == 1:
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                current_level = level
                
            # Deteksi bullet points
            elif line.startswith('* ') or line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                
            # Deteksi numbered lists
            elif re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line)
                p = doc.add_paragraph(text, style='List Number')
                
            # Cek apakah line ini adalah subheading (tanpa markdown #)
            elif len(line) < 50 and (line.isupper() or line.endswith(':')):
                doc.add_heading(line, level=2)
                
            # Paragraf biasa
            else:
                p = doc.add_paragraph(line)
                # Jika ini terlihat seperti paragraf panjang, justify teksnya
                if len(line) > 100:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Menambahkan footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Dokumen dihasilkan oleh DataWizard.App"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(path)
    print(f"[SUKSES] File Word disimpan ke {path}")
